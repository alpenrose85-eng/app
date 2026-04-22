import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import json
import io
import hashlib
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ModuleNotFoundError:
    DOCX_AVAILABLE = False

st.set_page_config(page_title="Расчёт остаточного ресурса змеевиков", layout="wide")
st.title("Определение остаточного ресурса змеевиков ВРЧ")

# --- Инициализация session_state ---
if 'test_data_input' not in st.session_state:
    st.session_state.test_data_input = []
if 'widget_prefix' not in st.session_state:
    st.session_state.widget_prefix = "default"
if 'steel_grade' not in st.session_state:
    st.session_state.steel_grade = "12Х1МФ"
if 'selected_param' not in st.session_state:
    st.session_state.selected_param = "Трунина"
if 'resource_calculations' not in st.session_state:
    st.session_state.resource_calculations = []

# --- Загрузка / сохранение проекта ---
st.sidebar.header("📁 Сохранить / загрузить проект")
uploaded_file = st.sidebar.file_uploader("Загрузите проект (.json)", type=["json"])
uploaded_excel = st.sidebar.file_uploader("Загрузите данные испытаний (.xlsx, .xls)", type=["xlsx", "xls"])
project_data = None

# Управление префиксом ключей для сброса кэша виджетов при загрузке
if uploaded_file is not None:
    try:
        project_data = json.load(uploaded_file)
        st.sidebar.success("✅ Проект загружен!")
        prefix_seed = json.dumps(project_data, sort_keys=True, ensure_ascii=False)
        st.session_state.widget_prefix = "loaded_" + str(hash(prefix_seed))[:12]
    except Exception as e:
        st.sidebar.error(f"❌ Ошибка при загрузке: {e}")
        st.session_state.widget_prefix = "default"
else:
    pass

# --- Загрузка данных из Excel ---
if uploaded_excel is not None:
    try:
        excel_bytes = uploaded_excel.getvalue()
        
        try:
            excel_data = pd.read_excel(io.BytesIO(excel_bytes), engine='openpyxl')
        except Exception:
            try:
                excel_data = pd.read_excel(io.BytesIO(excel_bytes), engine='xlrd')
            except Exception:
                excel_data = pd.read_excel(io.BytesIO(excel_bytes))
        
        # Проверяем наличие необходимых столбцов
        required_columns = ['Образец', 'sigma_MPa', 'T_C', 'tau_h']
        optional_columns = ['Группа_аппроксимации', 'group_approx', 'Группа']
        
        # Находим столбец группы аппроксимации
        group_col = None
        for col_name in optional_columns:
            if col_name in excel_data.columns:
                group_col = col_name
                break
        
        missing_columns = [col for col in required_columns if col not in excel_data.columns]
        
        if missing_columns:
            st.sidebar.error(f"❌ В файле Excel отсутствуют необходимые столбцы: {missing_columns}")
            st.sidebar.info("📋 Нужные столбцы: Образец, sigma_MPa, T_C, tau_h")
        else:
            test_data_from_excel = []
            for _, row in excel_data.iterrows():
                # Определяем группу аппроксимации
                if group_col:
                    try:
                        group_val = int(row[group_col])
                    except:
                        group_val = 0
                else:
                    group_val = 0
                
                test_data_from_excel.append({
                    "Образец": str(row['Образец']),
                    "sigma_MPa": float(row['sigma_MPa']),
                    "T_C": float(row['T_C']),
                    "tau_h": float(row['tau_h']),
                    "Группа_аппроксимации": group_val
                })
            
            st.session_state.test_data_input = test_data_from_excel
            st.sidebar.success(f"✅ Загружено {len(test_data_from_excel)} испытаний из Excel")
            
            data_str = json.dumps(test_data_from_excel, sort_keys=True)
            hash_obj = hashlib.md5(data_str.encode()).hexdigest()[:12]
            st.session_state.widget_prefix = f"excel_{hash_obj}"
            
    except Exception as e:
        st.sidebar.error(f"❌ Ошибка при чтении Excel файла: {str(e)}")

# --- Загрузка параметров или установка значений по умолчанию ---
if project_data is not None:
    loaded_test_data = project_data.get("испытания", [])
    params = project_data.get("параметры_трубы", {})
    selected_param = project_data.get("выбранный_параметр", "Трунина")
    selected_steel = project_data.get("марка_стали", "12Х1МФ")
    C_trunin_val = project_data.get("коэффициент_C_trunin", 24.88)
    C_larson_val = project_data.get("коэффициент_C_larson", 20.0)
    series_name = project_data.get("название_серии", "Образцы")
    st.session_state.test_data_input = loaded_test_data.copy()
    st.session_state.steel_grade = selected_steel
    st.session_state.selected_param = selected_param
    if 'resource_calculations' in project_data:
        st.session_state.resource_calculations = project_data['resource_calculations']
else:
    params = {}
    selected_param = "Трунина"
    selected_steel = st.session_state.steel_grade
    series_name = "Образцы"
    if not st.session_state.test_data_input:
        st.session_state.test_data_input = [{
            "Образец": f"Обр.{i+1}", 
            "sigma_MPa": 120.0, 
            "T_C": 600.0, 
            "tau_h": 500.0,
            "Группа_аппроксимации": 0
        } for i in range(6)]

# --- Название серии испытаний ---
st.header("0. Название серии испытаний")
series_name = st.text_input("Введите название серии образцов", value=series_name)

# --- Выбор марки стали ---
st.header("1. Выберите марку стали")
steel_options = ["12Х1МФ", "12Х18Н12Т", "ДИ82"]
selected_steel = st.selectbox(
    "Марка стали",
    options=steel_options,
    index=steel_options.index(selected_steel) if selected_steel in steel_options else 0
)
st.session_state.steel_grade = selected_steel

# --- Выбор типа параметра долговечности ---
st.header("2. Выберите параметр долговечности")
param_options = ["Трунина", "Ларсона-Миллера"]
selected_param = st.selectbox(
    "Тип параметра",
    options=param_options,
    index=param_options.index(selected_param) if selected_param in param_options else 0
)
st.session_state.selected_param = selected_param

# --- Автоматическая установка коэффициентов в зависимости от марки стали ---
def set_default_coefficients(steel_grade: str, parameter: str) -> float:
    """Устанавливает коэффициенты по умолчанию в зависимости от марки стали и параметра."""
    coefficients = {
        "12Х1МФ": {
            "Трунина": 24.88,
            "Ларсона-Миллера": 20.0
        },
        "12Х18Н12Т": {
            "Трунина": 26.3,
            "Ларсона-Миллера": 20.0
        },
        "ДИ82": {
            "Трунина": 39.87,
            "Ларсона-Миллера": 30.0
        }
    }
    return coefficients.get(steel_grade, {}).get(parameter, 20.0)


def format_approximation_equation(a: float, b: float, selected_param: str) -> str:
    """Возвращает уравнение аппроксимации в понятном виде."""
    sign = "+" if b >= 0 else "-"
    return f"lg(σ) = {a:.4f} · P {sign} {abs(b):.4f}"


def create_word_report(series_name: str, df_tests: pd.DataFrame, selected_param: str, c_value: float) -> io.BytesIO:
    """Создает Word-отчет по испытаниям."""
    if not DOCX_AVAILABLE:
        raise ModuleNotFoundError("python-docx не установлен")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"Отчет по испытаниям — {series_name}")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run(f"Параметр долговечности: {selected_param}")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)

    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Название образца",
        "Напряжение, МПа",
        "Температура, °C",
        "Время, ч",
        "Расчетный параметр Трунина",
    ]

    for i, header in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(header)
        rr.bold = True
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(14)

    for _, row in df_tests.iterrows():
        T_K = float(row["T_C"]) + 273.15
        tau_h = float(row["tau_h"])
        if selected_param == "Трунина":
            p_value = T_K * (np.log10(tau_h) - 2 * np.log10(T_K) + c_value) * 1e-3
        else:
            p_value = T_K * (np.log10(tau_h) + c_value) * 1e-3

        values = [
            str(row["Образец"]),
            f"{float(row['sigma_MPa']):.2f}".replace('.', ','),
            f"{float(row['T_C']):.1f}".replace('.', ','),
            f"{float(row['tau_h']):.1f}".replace('.', ','),
            f"{p_value:.4f}".replace('.', ','),
        ]
        cells = table.add_row().cells
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(value)
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(14)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Начальные значения коэффициентов по умолчанию
if project_data is None:
    C_trunin_val = set_default_coefficients(selected_steel, "Трунина")
    C_larson_val = set_default_coefficients(selected_steel, "Ларсона-Миллера")

# Применяем коэффициенты по умолчанию при изменении марки стали или параметра
if 'prev_steel' not in st.session_state:
    st.session_state.prev_steel = selected_steel
if 'prev_param' not in st.session_state:
    st.session_state.prev_param = selected_param

if (st.session_state.prev_steel != selected_steel or 
    st.session_state.prev_param != selected_param):
    default_C = set_default_coefficients(selected_steel, selected_param)
    if selected_param == "Трунина":
        C_trunin_val = default_C
    else:
        C_larson_val = default_C
    
    st.session_state.prev_steel = selected_steel
    st.session_state.prev_param = selected_param

# --- Настройка количества испытаний ---
st.header("3. Настройка количества испытаний")
num_tests = st.slider(
    "Количество испытаний (образцов)",
    min_value=0,
    max_value=100,
    value=len(st.session_state.test_data_input),
    step=1
)

# --- Синхронизация session_state с num_tests ---
if len(st.session_state.test_data_input) != num_tests:
    current = st.session_state.test_data_input
    if num_tests > len(current):
        for i in range(len(current), num_tests):
            current.append({
                "Образец": f"Обр.{i+1}", 
                "sigma_MPa": 120.0, 
                "T_C": 600.0, 
                "tau_h": 500.0,
                "Группа_аппроксимации": 0
            })
    else:
        current = current[:num_tests]
    st.session_state.test_data_input = current

# --- Ввод данных испытаний с группами аппроксимации ---
st.header("4. Введите данные испытаний")
st.info("""
💡 **Группы аппроксимации:**
- **0** = только на графике (без аппроксимации)
- **1, 2, 3...** = разные блоки точек для аппроксимации
Для построения аппроксимации в группе должно быть минимум 2 точки
""")

if num_tests > 0:
    for i in range(num_tests):
        col1, col2, col3, col4, col5 = st.columns([2, 2, 2, 2, 1])
        with col1:
            sample = col1.text_input(
                f"Образец {i+1}",
                value=st.session_state.test_data_input[i]["Образец"],
                key=f"{st.session_state.widget_prefix}_sample_{i}"
            )
        with col2:
            sigma = col2.number_input(
                f"σ, МПа",
                value=float(st.session_state.test_data_input[i]["sigma_MPa"]),
                min_value=0.1,
                max_value=500.0,
                key=f"{st.session_state.widget_prefix}_sigma_{i}"
            )
        with col3:
            T_C = col3.number_input(
                f"T, °C",
                value=float(st.session_state.test_data_input[i]["T_C"]),
                min_value=100.0,
                max_value=1000.0,
                key=f"{st.session_state.widget_prefix}_T_{i}"
            )
        with col4:
            tau_h = col4.number_input(
                f"τ, ч",
                value=float(st.session_state.test_data_input[i]["tau_h"]),
                min_value=1.0,
                max_value=1e7,
                key=f"{st.session_state.widget_prefix}_tau_{i}"
            )
        with col5:
            group = col5.number_input(
                f"Группа",
                value=int(st.session_state.test_data_input[i]["Группа_аппроксимации"]),
                min_value=0,
                max_value=10,
                step=1,
                key=f"{st.session_state.widget_prefix}_group_{i}"
            )
        st.session_state.test_data_input[i] = {
            "Образец": sample,
            "sigma_MPa": sigma,
            "T_C": T_C,
            "tau_h": tau_h,
            "Группа_аппроксимации": group
        }
else:
    st.info("Нет данных испытаний. График будет построен только с кривой допускаемых напряжений.")

df_tests = pd.DataFrame(st.session_state.test_data_input) if st.session_state.test_data_input else pd.DataFrame()

# --- Ввод общих параметров трубы для графика ---
st.header("5. Введите общие параметры трубы для графика")
col1, col2 = st.columns(2)
with col1:
    s_nom_val = params.get("s_nom", 6.0)
    s_nom = st.number_input("Номинальная толщина стенки s_н, мм", value=float(s_nom_val), min_value=0.1, max_value=1000.0)
    
    s_min_val = params.get("s_min", 5.07)
    s_min = st.number_input("Текущая min толщина s_мин, мм", value=float(s_min_val), min_value=0.1, max_value=s_nom)
    
    s_max_val = params.get("s_max", 5.95)
    s_max = st.number_input("Текущая max толщина s_макс, мм", value=float(s_max_val), min_value=0.1, max_value=1000.0)
    
    tau_exp_val = params.get("tau_exp", 317259)
    tau_exp = st.number_input("Наработка τ_э, ч", value=int(tau_exp_val), min_value=1, max_value=5_000_000)
with col2:
    d_max_val = params.get("d_max", 19.90)
    d_max = st.number_input("Макс. внутр. диаметр d_макс, мм", value=float(d_max_val), min_value=0.1, max_value=1000.0)
    
    T_rab_C_val = params.get("T_rab_C", 517.0)
    T_rab_C = st.number_input("Рабочая температура T_раб, °C", value=float(T_rab_C_val), min_value=100.0, max_value=1000.0)
    
    p_MPa_val = params.get("p_MPa", 27.93)
    p_MPa = st.number_input("Давление пара p, МПа", value=float(p_MPa_val), min_value=0.1, max_value=100.0)
    
    k_zapas_val = params.get("k_zapas", 1.5)
    k_zapas = st.number_input("Коэффициент запаса k_зап", value=float(k_zapas_val), min_value=1.0, max_value=5.0)

# --- Настройка коэффициентов и графика ---
st.header("6. Дополнительные настройки")
col1, col2, col3 = st.columns(3)

with col1:
    if selected_param == "Трунина":
        default_C_value = set_default_coefficients(selected_steel, selected_param)
        C = st.number_input(
            "Коэффициент C в параметре Трунина",
            value=float(default_C_value),
            min_value=0.0,
            max_value=50.0,
            format="%.3f",
            help=f"По умолчанию для {selected_steel}: {default_C_value}"
        )
    else:
        default_C_value = set_default_coefficients(selected_steel, selected_param)
        C = st.number_input(
            "Коэффициент C в параметре Ларсона-Миллера",
            value=float(default_C_value),
            min_value=0.0,
            max_value=50.0,
            format="%.3f",
            help=f"По умолчанию для {selected_steel}: {default_C_value}"
        )

with col2:
    # Выбор групп для построения аппроксимации
    if len(df_tests) > 0:
        # Находим все уникальные группы (кроме 0)
        groups = sorted([g for g in df_tests["Группа_аппроксимации"].unique() if g > 0])
        if groups:
            group_options = {}
            for g in groups:
                count = len(df_tests[df_tests["Группа_аппроксимации"] == g])
                if count >= 2:
                    group_options[g] = f"Группа {g} ({count} точек, можно аппроксимировать)"
                else:
                    group_options[g] = f"Группа {g} ({count} точек, мало для аппроксимации)"
            
            selected_approx_groups = st.multiselect(
                "Выберите группы для построения аппроксимации:",
                options=list(group_options.keys()),
                format_func=lambda x: group_options[x],
                default=[]
            )
        else:
            st.info("Назначьте точкам группы аппроксимации (1, 2, 3...)")
            selected_approx_groups = []
    else:
        selected_approx_groups = []

with col3:
    fig_width_cm = st.slider("Ширина графика (см)", min_value=12, max_value=20, value=17, step=1)
    fig_width_in = fig_width_cm / 2.54
    fig_height_cm = st.slider("Высота графика (см)", min_value=8, max_value=15, value=10, step=1)
    fig_height_in = fig_height_cm / 2.54

# --- Управление несколькими расчетами остаточного ресурса ---
st.header("7. Управление расчетами остаточного ресурса")

# Кнопка для добавления нового расчета
if st.button("➕ Добавить новый расчет остаточного ресурса"):
    st.session_state.resource_calculations.append({
        'id': len(st.session_state.resource_calculations) + 1,
        'name': f'Расчет {len(st.session_state.resource_calculations) + 1}',
        'params': {
            's_nom': s_nom,
            's_min': s_min,
            's_max': s_max,
            'tau_exp': tau_exp,
            'd_max': d_max,
            'T_rab_C': T_rab_C,
            'p_MPa': p_MPa,
            'k_zapas': k_zapas
        },
        'selected_group': 0,
        'results': None
    })

# Отображение существующих расчетов
for idx, calc in enumerate(st.session_state.resource_calculations):
    with st.expander(f"📊 {calc['name']} (ID: {calc['id']})", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            calc['name'] = st.text_input(f"Название расчета", value=calc['name'], key=f"calc_name_{idx}")
            
            # Выбор группы для расчета
            if len(df_tests) > 0 and selected_approx_groups:
                calc['selected_group'] = st.selectbox(
                    f"Выберите группу аппроксимации для расчета:",
                    options=[0] + selected_approx_groups,
                    format_func=lambda x: f"Только график" if x == 0 else f"Группа {x}",
                    index=selected_approx_groups.index(calc['selected_group']) + 1 if calc['selected_group'] in selected_approx_groups else 0,
                    key=f"calc_group_{idx}"
                )
            else:
                calc['selected_group'] = 0
                st.warning("Для расчета выберите группы аппроксимации в разделе 6")
        
        with col2:
            # Параметры трубы для этого расчета
            st.subheader("Параметры трубы")
            calc['params']['s_nom'] = st.number_input(
                "s_н, мм", value=float(calc['params']['s_nom']), 
                min_value=0.1, max_value=1000.0, key=f"calc_s_nom_{idx}"
            )
            calc['params']['s_min'] = st.number_input(
                "s_мин, мм", value=float(calc['params']['s_min']), 
                min_value=0.1, max_value=calc['params']['s_nom'], key=f"calc_s_min_{idx}"
            )
            calc['params']['d_max'] = st.number_input(
                "d_макс, мм", value=float(calc['params']['d_max']), 
                min_value=0.1, max_value=1000.0, key=f"calc_d_max_{idx}"
            )
            calc['params']['T_rab_C'] = st.number_input(
                "T_раб, °C", value=float(calc['params']['T_rab_C']), 
                min_value=100.0, max_value=1000.0, key=f"calc_T_rab_C_{idx}"
            )
            calc['params']['p_MPa'] = st.number_input(
                "p, МПа", value=float(calc['params']['p_MPa']), 
                min_value=0.1, max_value=100.0, key=f"calc_p_MPa_{idx}"
            )
            calc['params']['k_zapas'] = st.number_input(
                "k_зап", value=float(calc['params']['k_zapas']), 
                min_value=1.0, max_value=5.0, key=f"calc_k_zapas_{idx}"
            )
        
        # Кнопки управления
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        with col_btn1:
            if st.button(f"🔄 Рассчитать", key=f"calc_button_{idx}"):
                # Флаг для расчета будет установлен позже
                pass
        with col_btn2:
            if st.button(f"🗑️ Удалить", key=f"delete_calc_{idx}"):
                st.session_state.resource_calculations.pop(idx)
                st.rerun()
        with col_btn3:
            if st.button(f"📋 Скопировать параметры", key=f"copy_calc_{idx}"):
                # Копирование текущих параметров в новый расчет
                st.session_state.resource_calculations.append({
                    'id': len(st.session_state.resource_calculations) + 1,
                    'name': f'Копия {calc["name"]}',
                    'params': calc['params'].copy(),
                    'selected_group': calc['selected_group'],
                    'results': None
                })
                st.rerun()

# --- Кнопка сохранения проекта ---
if st.sidebar.button("💾 Сохранить проект"):
    data_to_save = {
        "название_серии": series_name,
        "марка_стали": selected_steel,
        "испытания": st.session_state.test_data_input,
        "параметры_трубы": {
            "s_nom": s_nom,
            "s_min": s_min,
            "s_max": s_max,
            "tau_exp": tau_exp,
            "d_max": d_max,
            "T_rab_C": T_rab_C,
            "p_MPa": p_MPa,
            "k_zapas": k_zapas
        },
        "выбранный_параметр": selected_param,
        "выбранные_группы_аппроксимации": selected_approx_groups,
        "коэффициент_C_trunin": C if selected_param == "Трунина" else 24.88,
        "коэффициент_C_larson": C if selected_param == "Ларсона-Миллера" else 20.0,
        "resource_calculations": st.session_state.resource_calculations
    }
    json_str = json.dumps(data_to_save, indent=2, ensure_ascii=False)
    st.sidebar.download_button(
        label="📥 Скачать проект (.json)",
        data=json_str,
        file_name="проект_ресурса.json",
        mime="application/json"
    )

# --- Скачать Word-отчет ---
if len(st.session_state.test_data_input) > 0:
    if DOCX_AVAILABLE:
        try:
            df_word_report = pd.DataFrame(st.session_state.test_data_input)
            c_for_report = C_trunin_val if selected_param == "Трунина" else C_larson_val
            word_report = create_word_report(series_name, df_word_report, selected_param, c_for_report)
            st.sidebar.download_button(
                label="📄 Скачать отчет Word",
                data=word_report,
                file_name="otchet_dlitelnoi_prochnosti.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.sidebar.warning(f"Не удалось сформировать Word-отчет: {e}")
    else:
        st.sidebar.info("Word-отчет недоступен: в окружении не установлен пакет python-docx")

# --- Шаблон Excel для скачивания ---
if st.sidebar.button("📥 Скачать шаблон Excel"):
    template_data = {
        'Образец': ['Обр.1', 'Обр.2', 'Обр.3', 'Обр.4', 'Обр.5'],
        'sigma_MPa': [120.0, 130.0, 140.0, 125.0, 135.0],
        'T_C': [600.0, 610.0, 620.0, 605.0, 615.0],
        'tau_h': [500.0, 450.0, 400.0, 480.0, 430.0],
        'Группа_аппроксимации': [1, 1, 1, 2, 2]
    }
    template_df = pd.DataFrame(template_data)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        template_df.to_excel(writer, index=False, sheet_name='Данные испытаний')
    
    st.sidebar.download_button(
        label="Скачать шаблон (.xlsx)",
        data=output.getvalue(),
        file_name="шаблон_данных_испытаний.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# --- Функция для расчета остаточного ресурса ---
def calculate_residual_resource(params: Dict, approx: Dict, selected_param: str, C: float, 
                                steel_grade: str, s_nom: float) -> Tuple[Optional[float], Dict]:
    """Рассчитывает остаточный ресурс по заданным параметрам."""
    
    try:
        # Извлекаем параметры
        s_min = params['s_min']
        s_max = params['s_max']
        tau_exp = params['tau_exp']
        d_max = params['d_max']
        T_rab_C = params['T_rab_C']
        p_MPa = params['p_MPa']
        k_zapas = params['k_zapas']
        
        T_rab = T_rab_C + 273.15
        a = approx['a']
        b = approx['b']
        
        # Скорость коррозии
        if s_max > s_nom:
            v_corr = (s_max - s_min) / tau_exp
        else:
            v_corr = (s_nom - s_min) / tau_exp
        
        # Итерационный расчет
        tau_prognoz = 50000.0
        iteration_data = []
        
        for iter_num in range(100):
            # Будущая минимальная толщина
            s_min2 = s_min - v_corr * tau_prognoz
            if s_min2 <= 0:
                return None, {"error": "Толщина стенки станет ≤ 0"}
            
            # Напряжение для будущего состояния
            sigma_k2 = (p_MPa / 2) * (d_max / s_min2 + 1)
            sigma_rasch2 = k_zapas * sigma_k2
            
            # Параметр P из уравнения аппроксимации
            P_rab = (np.log10(sigma_rasch2) - b) / a
            
            # Время до разрушения
            if selected_param == "Трунина":
                log_tau_r = P_rab / T_rab * 1000 + 2 * np.log10(T_rab) - C
            else:
                log_tau_r = P_rab / T_rab * 1000 - C
            
            tau_r = 10**log_tau_r
            
            iteration_data.append({
                "Итерация": iter_num + 1,
                "τ_прогн, ч": round(tau_prognoz, 0),
                "τ_р, ч": round(tau_r, 0),
                "Разница, ч": round(tau_prognoz - tau_r, 0),
                "s_min2, мм": round(s_min2, 3),
                "σ_расч2, МПа": round(sigma_rasch2, 1)
            })
            
            if not np.isfinite(tau_r) or tau_r <= 0:
                return None, {"error": "Некорректное время до разрушения"}
            
            delta = tau_prognoz - tau_r
            if abs(delta) <= 200:
                # Сходимость достигнута
                return tau_prognoz, {
                    "iterations": iteration_data,
                    "final_tau_r": tau_r,
                    "final_s_min2": s_min2,
                    "final_sigma_rasch2": sigma_rasch2,
                    "v_corr": v_corr,
                    "delta": delta,
                    "converged": True
                }
            
            # Коррекция прогноза
            correction = delta * 0.5
            correction = np.clip(correction, -10000, 10000)
            tau_prognoz = tau_prognoz - correction
            
            if tau_prognoz <= 0:
                tau_prognoz = 1000
        
        # Не сошлось за 100 итераций
        return tau_prognoz, {
            "iterations": iteration_data,
            "final_tau_r": tau_r,
            "final_s_min2": s_min2,
            "final_sigma_rasch2": sigma_rasch2,
            "v_corr": v_corr,
            "delta": delta,
            "converged": False
        }
        
    except Exception as e:
        return None, {"error": str(e)}

# --- Основной расчет и построение графика ---
if st.button("🚀 Построить график и выполнить расчеты"):
    try:
        # --- 1. Расчет напряжений для фактического состояния (общий график) ---
        sigma_fact_graph = (p_MPa / 2) * (d_max / s_min + 1)
        sigma_rasch = k_zapas * sigma_fact_graph
        T_rab = T_rab_C + 273.15
        
        if selected_param == "Трунина":
            P_fact = T_rab * (np.log10(tau_exp) - 2 * np.log10(T_rab) + C) * 1e-3
        else:
            P_fact = T_rab * (np.log10(tau_exp) + C) * 1e-3
        
        # --- 2. Расчет для точек испытаний (если есть) ---
        if len(df_tests) > 0:
            df_tests["T_K"] = df_tests["T_C"] + 273.15
            
            if selected_param == "Трунина":
                df_tests["P"] = df_tests["T_K"] * (np.log10(df_tests["tau_h"]) - 2 * np.log10(df_tests["T_K"]) + C) * 1e-3
            else:
                df_tests["P"] = df_tests["T_K"] * (np.log10(df_tests["tau_h"]) + C) * 1e-3
        
        # --- 3. Построение кривой допускаемых напряжений ---
        sigma_vals = np.linspace(20, 150, 300)
        
        if selected_steel == "12Х1МФ":
            P_dop = (24956 - 2400 * np.log10(sigma_vals) - 10.9 * sigma_vals) * 1e-3
            steel_label = f"12Х1МФ (допускаемое снижение длительной прочности)"
        elif selected_steel == "12Х18Н12Т":
            P_dop = (30942 - 3762 * np.log10(sigma_vals) - 16.8 * sigma_vals) * 1e-3
            steel_label = f"12Х18Н12Т (допускаемое снижение длительной прочности)"
        elif selected_steel == "ДИ82":
            P_dop = (40086 - 2400 * np.log10(sigma_vals) - 19.4 * sigma_vals) * 1e-3
            steel_label = f"ДИ82 (допускаемое снижение длительной прочности)"
        else:
            P_dop = (24956 - 2400 * np.log10(sigma_vals) - 10.9 * sigma_vals) * 1e-3
            steel_label = f"Допускаемое снижение длительной прочности"
        
        # --- 4. Аппроксимации для выбранных групп ---
        approximations = {}  # Будем хранить параметры аппроксимаций
        
        if len(df_tests) > 0 and selected_approx_groups:
            # Цвета для линий аппроксимации
            line_colors = {
                1: 'blue',
                2: 'red',
                3: 'green',
                4: 'purple',
                5: 'orange',
                6: 'brown',
                7: 'pink',
                8: 'cyan',
                9: 'magenta',
                10: 'olive'
            }
            
            for group_num in selected_approx_groups:
                group_data = df_tests[df_tests["Группа_аппроксимации"] == group_num]
                
                if len(group_data) >= 2:
                    X = group_data["P"].values
                    y = np.log10(group_data["sigma_MPa"].values)
                    A = np.vstack([X, np.ones(len(X))]).T
                    try:
                        a, b = np.linalg.lstsq(A, y, rcond=None)[0]
                        R2 = 1 - np.sum((y - (a*X + b))**2) / np.sum((y - np.mean(y))**2)
                        
                        approximations[group_num] = {
                            'a': a,
                            'b': b,
                            'R2': R2,
                            'count': len(group_data),
                            'color': line_colors.get(group_num, 'gray'),
                            'equation': format_approximation_equation(a, b, selected_param)
                        }
                    except Exception as e:
                        st.warning(f"Не удалось построить аппроксимацию для группы {group_num}: {str(e)}")
                else:
                    st.warning(f"Для группы {group_num} недостаточно точек для аппроксимации (нужно минимум 2)")
        
        # --- 5. Построение графика ---
        fig, ax = plt.subplots(figsize=(fig_width_in, fig_height_in))
        
        # 1. Кривая допускаемых напряжений
        ax.plot(P_dop, sigma_vals, 'k-', label=steel_label, linewidth=2)
        
        # 2. Точки испытаний (если есть)
        if len(df_tests) > 0:
            # Цвета для разных групп точек
            group_colors = {
                0: 'gray',      # Точки без аппроксимации
                1: 'blue',      # Группа 1
                2: 'red',       # Группа 2
                3: 'green',     # Группа 3
                4: 'purple',    # Группа 4
                5: 'orange',    # Группа 5
                6: 'brown',     # Группа 6
                7: 'pink',      # Группа 7
                8: 'cyan',      # Группа 8
                9: 'magenta',   # Группа 9
                10: 'olive'     # Группа 10
            }
            
            # Отображаем точки по группам
            for group_num in sorted(df_tests["Группа_аппроксимации"].unique()):
                group_data = df_tests[df_tests["Группа_аппроксимации"] == group_num]
                color = group_colors.get(group_num, 'gray')
                
                if group_num == 0:
                    label = f'Точки без аппроксимации ({len(group_data)} шт.)'
                elif group_num in selected_approx_groups:
                    if group_num in approximations:
                        label = f'Группа {group_num} ({len(group_data)} точек, R²={approximations[group_num]["R2"]:.3f})'
                    else:
                        label = f'Группа {group_num} ({len(group_data)} точек)'
                else:
                    label = f'Группа {group_num} ({len(group_data)} точек, аппроксимация не строится)'
                
                ax.scatter(group_data["P"], group_data["sigma_MPa"], 
                          c=color, s=50, label=label, alpha=0.7)
        
        # 3. Аппроксимационные линии для выбранных групп
        for group_num, approx_data in approximations.items():
            if group_num in selected_approx_groups:
                # Строим линию аппроксимации
                P_appr = (np.log10(sigma_vals) - approx_data['b']) / approx_data['a']
                ax.plot(P_appr, sigma_vals, color=approx_data['color'], linestyle='--', 
                       linewidth=1.5, label=f'Аппроксимация Гр.{group_num} (R²={approx_data["R2"]:.3f})')
        
        # 4. Фактическое состояние трубы (общий график)
        ax.scatter(P_fact, sigma_fact_graph, c='green', s=120, marker='o',
                  edgecolors='black', linewidth=1.5, 
                  label=f'Факт: σ={sigma_fact_graph:.1f} МПа (без запаса)')
        
        ax.scatter(P_fact, sigma_rasch, c='red', s=120, marker='s',
                  edgecolors='black', linewidth=1.5,
                  label=f'Расч: σ={sigma_rasch:.1f} МПа (k={k_zapas})')
        
        ax.plot([P_fact, P_fact], [sigma_fact_graph, sigma_rasch], 
               'k--', linewidth=1, alpha=0.5)
        
        # Настройка графика
        ax.set_xlim(P_dop.min() - 0.1, P_dop.max() + 0.1)
        ax.set_ylim(20, 150)
        
        if selected_param == "Трунина":
            xlabel_text = f"Параметр Трунина $P = T \\cdot (\\log_{{10}}(\\tau) - 2\\log_{{10}}(T) + {C:.2f}) \\cdot 10^{{-3}}$"
        else:
            xlabel_text = f"Параметр Ларсона-Миллера $P = T \\cdot (\\log_{{10}}(\\tau) + {C:.2f}) \\cdot 10^{{-3}}$"
        
        ax.set_xlabel(xlabel_text, fontsize=10)
        ax.set_ylabel(r"$\sigma$, МПа", fontsize=11)
        ax.set_title(f"Длительная прочность стали {selected_steel} - {series_name}", fontsize=12, pad=15)
        
        # Легенда справа от графика
        ax.legend(fontsize=8, frameon=True, fancybox=True, 
                 shadow=True, framealpha=0.9, 
                 bbox_to_anchor=(1.05, 1), loc='upper left')
        
        ax.grid(True, alpha=0.3)
        plt.subplots_adjust(right=0.75)
        
        st.pyplot(fig, use_container_width=False)
        
        # --- 6. Уравнения аппроксимации ---
        if approximations:
            st.header("Уравнения аппроксимации")
            eq_rows = []
            for group_num, approx_data in approximations.items():
                eq_rows.append({
                    "Группа": group_num,
                    "Уравнение": approx_data["equation"],
                    "R²": round(approx_data["R2"], 4),
                    "Количество точек": approx_data["count"]
                })
            st.table(pd.DataFrame(eq_rows))

        # --- 7. Результаты расчетов для общего графика ---
        st.header("Результаты расчетов (общий график)")
        
        # Находим допускаемое напряжение для P_fact
        idx = np.argmin(np.abs(P_dop - P_fact))
        sigma_dop = sigma_vals[idx]
        
        # Расчет запасов
        margin_fact = (sigma_dop / sigma_fact_graph - 1) * 100
        margin_rasch = (sigma_dop / sigma_rasch - 1) * 100
        
        # Таблица с результатами
        results_data = {
            "Параметр": [
                "Давление p",
                "Макс. диаметр d_max",
                "Мин. толщина s_min",
                "Формула для σ_факт",
                "σ_факт (без запаса)",
                "Коэффициент запаса k_зап",
                "σ_расч (с запасом)",
                "Допускаемое напряжение σ_доп",
                "Запас по σ_факт",
                "Запас по σ_расч",
                "Рабочая температура T_раб",
                "Наработка τ_э",
                "Параметр P",
                "Марка стали",
                "Параметр долговечности",
                "Коэффициент C",
                "Количество групп аппроксимации"
            ],
            "Значение": [
                f"{p_MPa:.2f} МПа",
                f"{d_max:.2f} мм",
                f"{s_min:.3f} мм",
                f"σ = (p/2) × (d/s + 1)",
                f"{sigma_fact_graph:.1f} МПа",
                f"{k_zapas:.1f}",
                f"{sigma_rasch:.1f} МПа",
                f"{sigma_dop:.1f} МПа",
                f"{margin_fact:.1f}%" + (" ✅" if margin_fact > 0 else " ⚠️"),
                f"{margin_rasch:.1f}%" + (" ✅" if margin_rasch > 0 else " ⚠️"),
                f"{T_rab_C:.1f} °C ({T_rab:.1f} K)",
                f"{tau_exp:,} ч",
                f"{P_fact:.4f}",
                selected_steel,
                selected_param,
                f"{C:.2f}",
                f"{len(selected_approx_groups)}" if selected_approx_groups else "0"
            ]
        }
        
        results_df = pd.DataFrame(results_data)
        st.table(results_df)
        
        # --- 7. Расчеты остаточного ресурса для каждого индивидуального расчета ---
        if st.session_state.resource_calculations and approximations:
            st.header("Результаты расчетов остаточного ресурса")
            
            for idx, calc in enumerate(st.session_state.resource_calculations):
                if calc['selected_group'] > 0 and calc['selected_group'] in approximations:
                    approx = approximations[calc['selected_group']]
                    
                    st.subheader(f"{calc['name']} (Группа аппроксимации: {calc['selected_group']})")
                    
                    # Выполняем расчет
                    tau_prognoz, calc_results = calculate_residual_resource(
                        calc['params'], approx, selected_param, C, selected_steel, calc['params']['s_nom']
                    )
                    
                    if tau_prognoz is not None:
                        if calc_results.get('converged', False):
                            st.success(f"✅ **Остаточный ресурс: {tau_prognoz:,.0f} ч** ({tau_prognoz/8760:.1f} лет)")
                            
                            # Таблица результатов
                            forecast_data = {
                                "Параметр": [
                                    "Остаточный ресурс τ_прогн",
                                    "Время до разрушения τ_р",
                                    "Разница (τ_прогн - τ_р)",
                                    "Минимальная толщина после ресурса",
                                    "Напряжение после ресурса (с запасом)",
                                    "Скорость коррозии",
                                    "Коэффициент аппроксимации a",
                                    "Коэффициент аппроксимации b",
                                    "R² аппроксимации",
                                    "Количество точек в группе",
                                    "Коэффициент запаса",
                                    "Рабочая температура",
                                    "Давление"
                                ],
                                "Значение": [
                                    f"{tau_prognoz:,.0f} ч",
                                    f"{calc_results['final_tau_r']:,.0f} ч",
                                    f"{calc_results['delta']:.0f} ч",
                                    f"{calc_results['final_s_min2']:.3f} мм",
                                    f"{calc_results['final_sigma_rasch2']:.1f} МПа",
                                    f"{calc_results['v_corr']:.6f} мм/ч",
                                    f"{approx['a']:.4f}",
                                    f"{approx['b']:.4f}",
                                    f"{approx['R2']:.4f}",
                                    f"{approx['count']}",
                                    f"{calc['params']['k_zapas']:.1f}",
                                    f"{calc['params']['T_rab_C']:.1f} °C",
                                    f"{calc['params']['p_MPa']:.2f} МПа"
                                ]
                            }
                            
                            forecast_df = pd.DataFrame(forecast_data)
                            st.table(forecast_df)
                            
                            # Итерации (если нужно)
                            if st.checkbox(f"Показать процесс итераций для {calc['name']}", key=f"show_iters_{idx}"):
                                st.table(pd.DataFrame(calc_results['iterations']))
                        else:
                            st.warning(f"⚠️ Не удалось достичь полной сходимости. Последнее значение: {tau_prognoz:,.0f} ч")
                            st.info(f"Разница между прогнозом и временем до разрушения: {calc_results['delta']:.0f} ч")
                            
                            if st.checkbox(f"Показать процесс итераций для {calc['name']}", key=f"show_iters_{idx}"):
                                st.table(pd.DataFrame(calc_results['iterations']))
                    else:
                        st.error(f"❌ Ошибка в расчете: {calc_results.get('error', 'Неизвестная ошибка')}")
                elif calc['selected_group'] == 0:
                    st.info(f"Для расчета {calc['name']} не выбрана группа аппроксимации")
                else:
                    st.warning(f"Для расчета {calc['name']} выбранная группа {calc['selected_group']} не имеет аппроксимации")
        else:
            if not approximations:
                st.info("""
                **Для расчета остаточного ресурса:**
                1. Назначьте точкам номера групп аппроксимации (1, 2, 3...)
                2. Выберите группы для построения аппроксимации в разделе "Дополнительные настройки"
                3. В группе должно быть минимум 2 точки для построения аппроксимации
                4. Добавьте расчеты остаточного ресурса в разделе 7
                """)
            elif not st.session_state.resource_calculations:
                st.info("Добавьте расчеты остаточного ресурса в разделе 7 для получения результатов")
            
    except Exception as e:
        st.error(f"Ошибка: {str(e)[:200]}")
        import traceback
        st.text(traceback.format_exc())
